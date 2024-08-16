from django.shortcuts import render,redirect
from django.contrib import messages
from .forms import PDFUploadForm
from .models import PDFUpload
from .models import templates_obj
import threading
from .pdftodb import main as startpdfcon
from .pdftodbnewformat import main_new as newstartpdfcon
import sqlite3
from django.http import HttpResponse
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Alignment
import string,json
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn



BASE_DIR = Path(__file__).resolve().parent.parent
import os


def uploadpdf_record(request,pk):
    if request.user.is_authenticated:
        customer_record = PDFUpload.objects.get(id=pk)
        return render(request,'uploadpdf/uploadpdf_record.html',{'customer_record':customer_record})
    else:
        messages.success(request,"You Must Be Logged In To View This Page.....")
        return redirect('login')
    
def delete_uploadpdf(request,pk):
    if request.user.is_authenticated:
        delete_it = PDFUpload.objects.get(id=pk)
        file_path = delete_it.pdf_file.path  # Assuming 'file' is the field name in PDFUpload model
        
        # Check if file exists and then delete it
        if os.path.exists(file_path):
            os.remove(file_path)
        delete_it.delete()
        messages.success(request,"Records Deleted Successfully...!")
        return redirect('uploadpdf')
    else:
        messages.success(request,"You Must Be Logged In To View This Page.....")
        return redirect('login')


def add_uploadpdf(request):
    if request.user.is_authenticated:
        form = PDFUploadForm(request.POST, request.FILES)
        if request.method == 'POST':
            
            if form.is_valid():
                cur_ins = form.save(commit=False)
                cur_ins.created_by = request.user.get_full_name()
                dynamic_fields = request.POST.getlist('dynamic_field')
                dicta={}
                alpha = list(string.ascii_uppercase)
                i = 0
                print(dynamic_fields)
                for field_data in dynamic_fields:
                    field_data = re.split(r'[,\s\t\n]+',field_data)
                    field_data = list(filter(None, field_data))
                    #print(str(alpha[i]) + " "+ str(field_data))    #Handle each field data (e.g., save to the database)
                    dicta[alpha[i]]=field_data
                    i +=1
                print(dicta)
                print(type(dicta))
                cur_ins.div_PRN = dicta
                cur_ins.save()
                messages.success(request,"PDF Uploaded Successfully...!")
                # t2 = threading.Thread(target=startpdfcon, name='t2')
                # t2.start()
                if cur_ins.new_pdf:
                    newstartpdfcon()
                else:
                    startpdfcon()
                return redirect('uploadpdf')
                
        else:
            form = PDFUploadForm()
        return render(request, 'uploadpdf/add_uploadpdf.html', {'form': form})
    else:
        messages.success(request,"You Must Be Logged In To View This Page.....")
        return redirect('login')
    
def update_uploadpdf(request,pk):
    if request.user.is_authenticated:
        record = PDFUpload.objects.get(id=pk)
        if request.method == 'POST':
            form = PDFUploadForm(request.POST, request.FILES, instance=record)
            
            # Set the static value for created_by
            #form.instance.created_by = request.user.get_full_name()
            form.instance.status     = ""
            dynamic_fields = request.POST.getlist('dynamic_field')
            dicta={}
            alpha = list(string.ascii_uppercase)
            i = 0
            print(dynamic_fields)
            for field_data in dynamic_fields:
                field_data = re.split(r'[,\s\t\n]+',field_data)
                field_data = list(filter(None, field_data))
                #print(str(alpha[i]) + " "+ str(field_data))    #Handle each field data (e.g., save to the database)
                dicta[alpha[i]]=field_data
                i +=1
            #print(dicta)
            #print(type(dicta))
            form.instance.div_PRN = dicta
            
            
            # Assign dynamic fields data to the form instance
            #print(dicta)
            form.instance.div_PRN = dicta
            
            if form.is_valid():
                form.save()
                messages.success(request, "PDF Updated Successfully!")
                #t2 = threading.Thread(target=startpdfcon, name='t2')
                #t2.start()
                print(form.instance.new_pdf)
                if form.instance.new_pdf:
                    newstartpdfcon()
                else:
                    startpdfcon()
                return redirect('uploadpdf')  # Replace with your success URL name
            
        else:
            form = PDFUploadForm(instance=record)
            
        return render(request, 'uploadpdf/add_uploadpdf.html', {'form': form})
    else:
        messages.success(request,"You Must Be Logged In To View This Page.....")
        return redirect('login')
    
def uploadpdf(request):
    
    if request.user.is_authenticated:
        pdfs = PDFUpload.objects.all().order_by("-created_at")
        return render(request,'uploadpdf.html',{'pdfs':pdfs})

    else:
        messages.success(request,"You Must Be Logged In To View This Page.....")
        return redirect('login')
  
def downloadcsv(request):
    if request.user.is_authenticated:
        if request.method == 'POST':
            tb_nm = request.POST['Name']
            tb_nmlist = str(tb_nm).split('_')
            new = False
            if tb_nmlist[0]=="NEW":
                new = True
                tempftb = tb_nmlist[1]+"_"+tb_nmlist[2]
            else:
                tempftb = tb_nmlist[0]+"_"+tb_nmlist[1]
            temp = templates_obj.objects.get(Name=tempftb)
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = 'attachment; filename="'+tb_nm+'.xlsx"'
            
            BE = False
            if (tb_nm.split('_'))[0]=="BE":
                BE = True
            
            connection = sqlite3.connect(BASE_DIR / 'data.db')
            connection.execute('PRAGMA journal_mode=WAL;')

            cursor = connection.cursor()
            wb = Workbook()
            #divs = ['A','B']
            cursor.execute(f"SELECT DISTINCT div FROM {tb_nm}")
            divs = [row[0] for row in cursor.fetchall()]
            if None in divs:
                divs.remove(None)
            divs.sort()
            divs += ['ALL',None] 
            #print(divs)
            for div in divs:
                
                ws = wb.create_sheet()
                
                
                
                
                if div == None:
                    ws.title="Backlog"
                    ws.append([])
                    ws.append([])
                    cursor.execute(f"SELECT * FROM {tb_nm} WHERE div is NULL ")
                elif div == 'ALL':
                    ws.title="ALL"
                    ws.append([])
                    ws.append([])
                    cursor.execute(f"SELECT * FROM {tb_nm} ")
                else:
                    ws.title=f"Div {div}"
                    ws.append([])
                    ws.append([])
                    cursor.execute(f"SELECT * FROM {tb_nm} WHERE div='{div}'")

                rows = cursor.fetchall()

                column_names = [description[0].upper() for description in cursor.description]
                if BE:
                    ind = column_names.index("CGPA")+3
                else:
                    ind = column_names.index("SGPA")+3
                    
                ws.append(column_names)

                for row in rows:
                    ws.append(row)#file.write(','.join(map(str, row)) + '\n')
                for key, value in temp.key_value_pairs.items():
                    #print(key,value[0],value.count(True))
                    if value[1]==True:
                        ws.merge_cells(start_row=1,end_row=2,start_column=ind,end_column=ind+value.count(True)+1)
                        ws.cell(row=1,column=ind).value=value[0]
                        ws.cell(row=1,column=ind).alignment = Alignment(horizontal='center',vertical='center')
                        ind= ind+value.count(True)+2
                    else:
                        ws.merge_cells(start_row=1,end_row=2,start_column=ind,end_column=ind+value.count(True)-1)
                        ws.cell(row=1,column=ind).value=value[0]
                        ws.cell(row=1,column=ind).alignment = Alignment(horizontal='center',vertical='center')
                        ind= ind+value.count(True)
                        
            default_sheet = wb['Sheet']
            wb.remove(default_sheet)
            connection.close()
            wb.save(response)
            return response

    else:
        messages.success(request,"You Must Be Logged In To View This Page.....")
        return redirect('login') 
    
    
def dwdresaysdocx(request):
    if request.user.is_authenticated:
        if request.method == 'POST':
            tb_nm = request.POST.get('Name', '')
            if not tb_nm:
                messages.error(request, "Name parameter is missing.")
                return redirect('uploadpdf')

            tb_nmlist = str(tb_nm).split('_')
            if len(tb_nmlist) < 2:
                messages.error(request, "Invalid table name format.")
                return redirect('uploadpdf')

            new = False
            if tb_nmlist[0]=="NEW":
                new = True
                tempftb = tb_nmlist[1]+"_"+tb_nmlist[2]
            else:
                tempftb = tb_nmlist[0]+"_"+tb_nmlist[1]
            try:
                temp = templates_obj.objects.get(Name=tempftb)
            except templates_obj.DoesNotExist:
                messages.error(request, "Template does not exist.")
                return redirect('uploadpdf')

            BE = tb_nmlist[0] == "BE"

        
            connection = sqlite3.connect(BASE_DIR / 'data.db')
            connection.execute('PRAGMA journal_mode=WAL;')
            cursor = connection.cursor()

            cursor.execute(f"SELECT DISTINCT div FROM {tb_nm}")
            divs = [row[0] for row in cursor.fetchall() if row[0] is not None and row[0] != ""]
            divs.sort()
            divs.insert(0, "ALL")
            if divs==["ALL"]:
                content="You haven't Entered PRN Division Wise"
                response = HttpResponse(content, content_type='text/plain')
                response['Content-Disposition'] = f'attachment; filename="{tb_nm}.txt"'
                return response

            document = Document()
            document.add_heading(f'Result Anaysis Report for {tb_nm}', level=1)

            for MYDIV in divs:
                # Add a page break before adding data for each division except the first one
                if MYDIV != "ALL":
                    document.add_page_break()

                document.add_heading("ALL Divisions" if MYDIV == "ALL" else f"{MYDIV} Division", level=2)

                query_condition = "div IS NOT NULL AND div != ''" if MYDIV == "ALL" else f"div='{MYDIV}'"

                q1 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition};"
                cursor.execute(q1)
                res1 = cursor.fetchone()[0] or 0
                if new:
                    q2 = f"SELECT count(name) FROM {tb_nm} WHERE SGPA!='-----'AND {query_condition};"
                else:
                    q2 = f"SELECT count(name) FROM {tb_nm} WHERE SGPA!='--'AND {query_condition};"
                

                cursor.execute(q2)
                res2 = cursor.fetchone()[0] or 0

                table = document.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'APPEARED'
                hdr_cells[1].text = 'TOTAL PASSING %'
                row_cells = table.add_row().cells
                row_cells[0].text = str(res1)
                row_cells[1].text = f"{(res2 / res1) * 100:.2f}" if res1 > 0 else '0.00'

                document.add_paragraph("\nTop 5 SGPA Students:")
                table = document.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Name'
                hdr_cells[1].text = 'SGPA'
                
                if new:
                    q = f"SELECT name, SGPA FROM {tb_nm} WHERE SGPA != '-----' AND {query_condition}   ORDER BY SGPA DESC LIMIT 5;"
                else:
                    q = f"SELECT name, SGPA FROM {tb_nm} WHERE SGPA != '--' AND {query_condition}   ORDER BY SGPA DESC LIMIT 5;"
                cursor.execute(q)
                result = cursor.fetchall()
                for res in result:
                    row_cells = table.add_row().cells
                    row_cells[0].text = res[0]
                    row_cells[1].text = str(res[1]) if res[1] is not None else 'N/A'

                document.add_paragraph("\nSubject Statistics:")
                table = document.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Sub Name'
                hdr_cells[1].text = 'PRESENT'
                hdr_cells[2].text = 'ABSENT'
                hdr_cells[3].text = 'PASS'
                hdr_cells[4].text = 'FAIL'
                hdr_cells[5].text = 'Result in %'

                for key, values in temp.key_value_pairs.items():
                    if values[1]:
                        if new:
                            query0 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_end != 'AAA' AND {query_condition};"
                            query1 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_end = 'AAA' AND {query_condition};"
                        else:
                            query0 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_end != 'AB' AND {query_condition};"
                            query1 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_end = 'AB' AND {query_condition};"
                            
                        query2 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_tot >= 40 AND {query_condition};"
                        query3 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_tot < 40 AND {query_condition};"

                        cursor.execute(query0)
                        result0 = cursor.fetchone()[0] or 0

                        cursor.execute(query1)
                        result1 = cursor.fetchone()[0] or 0

                        cursor.execute(query2)
                        result2 = cursor.fetchone()[0] or 0

                        cursor.execute(query3)
                        result3 = cursor.fetchone()[0] or 0

                        row_cells = table.add_row().cells
                        row_cells[0].text = values[0]
                        row_cells[1].text = str(result0)
                        row_cells[2].text = str(result1)
                        row_cells[3].text = str(result2)
                        row_cells[4].text = str(result3)
                        row_cells[5].text = f"{(result2 / result0) * 100:.2f}" + "%" if result0 > 0 else '0.00'
                document.add_page_break()
                document.add_paragraph("\nTop Marks in Each Subject:")
                table = document.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Sub Name'
                hdr_cells[1].text = 'Name'
                hdr_cells[2].text = 'Marks'

                for key, values in temp.key_value_pairs.items():
                    if values[1]:
                        if new:
                            query1 = f"SELECT name, _{key}_tot FROM {tb_nm} WHERE _{key}_end != 'AAA' AND {query_condition} ORDER BY _{key}_tot DESC LIMIT 1;"
                        else:
                            query1 = f"SELECT name, _{key}_tot FROM {tb_nm} WHERE _{key}_end != 'AB'  AND {query_condition} ORDER BY _{key}_tot DESC LIMIT 1;"
                        cursor.execute(query1)
                        result1 = cursor.fetchone()
                        if result1:
                            row_cells = table.add_row().cells
                            row_cells[0].text = values[0]
                            row_cells[1].text = result1[0]
                            row_cells[2].text = str(result1[1])
                        else:
                            row_cells = table.add_row().cells
                            row_cells[0].text = values[0]
                            row_cells[1].text = 'N/A'
                            row_cells[2].text = 'N/A'

                document.add_paragraph("\nSGPA Classification:")
                table = document.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Distinction'
                hdr_cells[1].text = 'First'
                hdr_cells[2].text = 'Higher Second'
                hdr_cells[3].text = 'Fail'
                if new:
                    query1 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '-----' AND SGPA >= 7.75;"
                    query2 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '-----' AND SGPA < 7.75 AND SGPA >= 6.75;"
                    query3 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '-----' AND SGPA < 6.75;"
                    query4 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA = '-----';"
                else:
                    query1 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '--'  AND SGPA >= 7.75;"
                    query2 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '--'  AND SGPA < 7.75 AND SGPA >= 6.75;"
                    query3 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '--'  AND SGPA < 6.75;"
                    query4 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA = '--' ;"


                cursor.execute(query1)
                result1 = cursor.fetchone()[0] or 0

                cursor.execute(query2)
                result2 = cursor.fetchone()[0] or 0

                cursor.execute(query3)
                result3 = cursor.fetchone()[0] or 0

                cursor.execute(query4)
                result4 = cursor.fetchone()[0] or 0

                row_cells = table.add_row().cells
                row_cells[0].text = str(result1)
                row_cells[1].text = str(result2)
                row_cells[2].text = str(result3)
                row_cells[3].text = str(result4)
                
                row_cells = table.add_row().cells
                row_cells[0].text = str(f"{(result1/res1)*100:.2f}") + "%"
                row_cells[1].text = str(f"{(result2/res1)*100:.2f}") + "%"
                row_cells[2].text = str(f"{(result3/res1)*100:.2f}") + "%"
                row_cells[3].text = str(f"{(result4/res1)*100:.2f}") + "%"

            connection.close()

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{tb_nm}.docx"'
            document.save(response)
            return response

            # except sqlite3.Error as e:
            #     messages.error(request, f"Database error: {e}")
            #     return redirect('uploadpdf')
            # except Exception as e:
            #     messages.error(request, f"An unexpected error occurred: {e}")
            #     return redirect('uploadpdf')

    else:
        messages.success(request, "You Must Be Logged In To View This Page.....")
        return redirect('login')
    
    
    
# def dwdresays(request):
#     if request.user.is_authenticated:
#         if request.method == 'POST':
#             tb_nm = request.POST['Name']
#             tb_nmlist = str(tb_nm).split('_')
#             tempftb = tb_nmlist[0] + "_" + tb_nmlist[1]
#             temp = templates_obj.objects.get(Name=tempftb)
            
#             BE = False
#             if tb_nmlist[0] == "BE":
#                 BE = True
            
#             connection = sqlite3.connect(BASE_DIR / 'data.db')
#             connection.execute('PRAGMA journal_mode=WAL;')

#             cursor = connection.cursor()
            
#             cursor.execute(f"SELECT DISTINCT div FROM {tb_nm}")
#             divs = [row[0] for row in cursor.fetchall()]
#             if None in divs:
#                 divs.remove(None)
#             if "" in divs:
#                 divs.remove("")
#             divs.sort()
#             divs.insert(0, "ALL")
#             content = ""
#             if divs==["ALL"]:
#                 content="You haven't Entered PRN Division Wise"
#                 response = HttpResponse(content, content_type='text/plain')
#                 response['Content-Disposition'] = f'attachment; filename="{tb_nm}.txt"'
#                 return response
#             for MYDIV in divs:
#                 if MYDIV == "ALL":
#                     query_condition = "div IS NOT NULL AND div != ''"
#                 else:
#                     query_condition = f"div='{MYDIV}'"
                
#                 content += ("-------------------------------------------------------------------------------------------------------------\n")
#                 if MYDIV == "ALL":
#                     content += ("\n\nALL Divisions\n\n")
#                 else:
#                     content += (f"\n\n{MYDIV} Division\n\n")
                    
#                 q1 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition};"
#                 cursor.execute(q1)
#                 res1 = cursor.fetchone()[0]
                
#                 q2 = f"SELECT count(name) FROM {tb_nm} WHERE SGPA!='--' AND {query_condition};"
#                 cursor.execute(q2)
#                 res2 = cursor.fetchone()[0]
                
#                 content+=f"{'APPEARED':<15} {'TOTAL PASSING %':<8}\n"
#                 content+=f"{res1:<15} {(res2/res1)*100:<8.2f}\n\n"
                
                
#                 q = f"SELECT name, SGPA FROM {tb_nm} WHERE {query_condition} AND SGPA != '--' ORDER BY SGPA DESC LIMIT 5;"
#                 cursor.execute(q)
#                 result = cursor.fetchall()
                
#                 content += (f"{'Name':<30} {'SGPA':<8} \n")
#                 for res in result:
#                     content += (f"{res[0]:<30} {res[1]:<8} \n")

#                 content += ("\n\n")
                
#                 content += (f"{'Sub Name':<40} {'PRESENT':<8}  {'ABSENT':<8} {'PASS':<8} {'FAIL':<8} {'Result in %':<8}\n")
#                 for key, values in temp.key_value_pairs.items():
#                     if values[1]:
#                         query0 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_end != 'AB' AND {query_condition};"
#                         query1 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_end = 'AB' AND {query_condition};"
#                         query2 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_tot >= 40 AND {query_condition};"
#                         query3 = f"SELECT count(name) FROM {tb_nm} WHERE _{key}_tot < 40 AND {query_condition};"

#                         cursor.execute(query0)
#                         result0 = cursor.fetchone()[0]
                    
#                         cursor.execute(query1)
#                         result1 = cursor.fetchone()[0]

#                         cursor.execute(query2)
#                         result2 = cursor.fetchone()[0]

#                         cursor.execute(query3)
#                         result3 = cursor.fetchone()[0]

#                         content += (f"{values[0]:<40} {result0:<8}  {result1:<8} {result2:<8} {result3:<8} {((result2 / result0) * 100):<8.2f}\n")
                        
#                 content += ("\n\n")
                
#                 content += (f"{'Sub Name':<40} {'Name':<50}  {'Marks':<8}\n")

#                 for key, values in temp.key_value_pairs.items():
#                     if values[1]:
#                         query1 = f"SELECT name, _{key}_tot FROM {tb_nm} WHERE _{key}_end != 'AB' AND {query_condition} ORDER BY _{key}_tot DESC LIMIT 1;"

#                         cursor.execute(query1)
#                         result1 = cursor.fetchone()

#                         content += (f"{values[0]:<40} {result1[0]:<50}  {result1[1]:<8}\n")
                
#                 content += ("\n\n")

#                 query1 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '--' AND SGPA >= 7.75;"
#                 query2 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '--' AND SGPA < 7.75 AND SGPA >= 6.75;"
#                 query3 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA != '--' AND SGPA < 6.75;"
#                 query4 = f"SELECT count(name) FROM {tb_nm} WHERE {query_condition} AND SGPA = '--';"

#                 cursor.execute(query1)
#                 result1 = cursor.fetchone()[0]     
                    
#                 cursor.execute(query2)
#                 result2 = cursor.fetchone()[0]

#                 cursor.execute(query3)
#                 result3 = cursor.fetchone()[0]

#                 cursor.execute(query4)
#                 result4 = cursor.fetchone()[0]
                
#                 content += (f"{'Distinction':<15} {'First':<15} {'Higher Second':<15} {'Fail':<15}\n")
#                 content += (f"{result1:<15} {result2:<15} {result3:<15} {result4:<15}\n")

#             connection.close()
#             response = HttpResponse(content, content_type='text/plain')
#             response['Content-Disposition'] = f'attachment; filename="{tb_nm}.txt"'
#             return response

#     else:
#         messages.success(request, "You Must Be Logged In To View This Page.....")
#         return redirect('login')
